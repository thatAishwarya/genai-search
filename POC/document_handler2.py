import os
import PyPDF2
import docx
from typing import List
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentHandler():
    def __init__(self, path: str, switch: bool):

        if switch:
            self.dp = path
            self.texts = []

            for root, _, files in os.walk(self.dp):
                for file in files:
                    file_path = os.path.join(root, file)
                    try:
                        text = self.extract_text(file_path)
                        self.texts.append(text)
                    except Exception as e:
                        logger.error(f"Error processing file {file_path}: {e}")
        else:
            self.fp = path
            self.result = self.extract_text(self.fp)

    def extract_text(self, file_path: str) -> str:
        if file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif file_path.endswith('.txt'):
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError("Unsupported file format")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        self.text = ""
        with open(pdf_path, 'rb') as file:
            self.reader = PyPDF2.PdfReader(file)
            for page in self.reader.pages:
                self.text += page.extract_text()
        return self.text

    def extract_text_from_docx(self, docx_path: str) -> str:
        self.doc = docx.Document(docx_path)
        self.text = "\n".join([para.text for para in self.doc.paragraphs])
        return self.text

    def extract_text_from_txt(self, txt_path: str) -> str:
        with open(txt_path, 'r', encoding='utf-8') as self.file:
            return self.file.read()

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_txt(txt_path: str) -> str:
    with open(txt_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text(file_path: str) -> str:
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError("Unsupported file format")

def read_documents_from_directory(directory_path: str) -> List[str]:
    texts = []
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                text = extract_text(file_path)
                texts.append(text)
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")
    return texts
